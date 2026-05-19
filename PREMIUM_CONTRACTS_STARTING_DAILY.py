# -*- coding: utf-8 -*-
"""
Created on Mon Dec  4 16:38:51 2023

@author: Leane Fraser
"""
import psycopg2
import psycopg2.extras
import win32com.client as win32             # pip install pypiwin32 (if not already installed)
import pandas as pd
import os
import sys
from lxml import etree
from datetime import datetime, timedelta
import docx #pip install python-docx

print("-------------------------------------------------------")
print("|                                                     |")
print("|          PREMIUM CONTRACTS STARTNG TODAY            |")
print("|                                                     |")
print("-------------------------------------------------------")

flyn_positions_host = "your-rds-endpoint.eu-west-3.rds.amazonaws.com"
read_db_username = "your_read_only_user"
db_password = "your_secure_password"
positions_db_name = "YourDatabaseName"

# Establish connection using generic variables
conn = psycopg2.connect(
    host=flyn_positions_host,
    database=positions_db_name,
    user=read_db_username,
    password=db_password,
)

def title():

        run.bold = True
        run.italic = True
        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(12)

def main_body_text():

        run.font.name = 'Arial'
        run.font.size = docx.shared.Pt(12)

now = datetime.now()
today = now.strftime("%Y-%m-%d")

tomorrow = now + timedelta(days = 1)
tomorrow = tomorrow.strftime("%Y-%m-%d")

year = datetime.now().strftime("%Y")
logo_absolute = str("C:\project_data\templates\logo.jpg")

sql = f"""SELECT c.producttype,TYPE,CAST(c.customerid AS VARCHAR),c.timezone,c.id,c.name,c.mmsi,c.yachttype,c.yachtname,c.predeliverydateutc,c.startdateutc,
c.enddateutc,c.postdeliverydateutc,c.captainname,c.captaincellular,c.captainemail
FROM "FLYNContractsData"."FLYNContracts" AS c
WHERE startdateutc > '{today}' AND startdateutc < '{tomorrow}'
AND status IN ('ACTIVE','DELIVERED','CLOSED','PLANNED') AND producttype = 'PREMIUM' OR predeliverydateutc > '{today}' AND predeliverydateutc < '{tomorrow}'
AND status IN ('ACTIVE','DELIVERED','CLOSED','PLANNED') AND producttype = 'PREMIUM' ORDER BY producttype, customerid ;"""

cursor.execute(sql)
premium = cursor.fetchall()

for contract in range(0,len(premium)):
    timezone = premium[contract]['timezone']
    if timezone == 'CEST':
        try:
            premium[contract]['startdateutc'] = premium[contract]['startdateutc'] + timedelta(hours=2)
            premium[contract]['enddateutc'] = premium[contract]['enddateutc'] + timedelta(hours=2)
            premium[contract]['predeliverydateutc'] = premium[contract]['predeliverydateutc'] + timedelta(hours=2)
            premium[contract]['postdeliverydateutc'] = premium[contract]['postdeliverydateutc'] + timedelta(hours=2)
        except:
            print(" ")

    if timezone == 'CET':
        try:
            premium[contract]['startdateutc'] = premium[contract]['startdateutc'] + timedelta(hours=1)
            premium[contract]['enddateutc'] = premium[contract]['enddateutc'] + timedelta(hours=1)
            premium[contract]['predeliverydateutc'] = premium[contract]['predeliverydateutc'] + timedelta(hours=1)
            premium[contract]['postdeliverydateutc'] = premium[contract]['postdeliverydateutc'] + timedelta(hours=1)
        except:
            print(" ")

    if timezone == 'AST':
        try:
            premium[contract]['startdateutc'] = premium[contract]['startdateutc'] - timedelta(hours=4)
            premium[contract]['enddateutc'] = premium[contract]['enddateutc'] - timedelta(hours=4)
            premium[contract]['predeliverydateutc'] = premium[contract]['predeliverydateutc'] - timedelta(hours=4)
            premium[contract]['postdeliverydateutc'] = premium[contract]['postdeliverydateutc'] - timedelta(hours=4)
        except:
            print(" ")

df = pd.DataFrame.from_dict(premium)
today_hrs =f'{today} 00:00:00'

try:
    df2 = df.drop(df[df['predeliverydateutc'] < today_hrs].index)
    df2.reset_index(drop=True, inplace=True)
    nb_ctrs_end_td = len(df2)
except:
    print(f"No premium contracts starting today ({today})")
    sys.exit()

no_pr = len(df2)
txt_pre = int(no_pr)

print(f"No PREMIUM contracts starting today. ({today})\n") if len(df2) <1 else print(f"{txt_pre} Premium contract(s) starting today.\n")
nb_c = len(df2)

for index, row in df2.iterrows():

    doc = docx.Document()
    id = f"{row['id']}"
    yachttype = f"{row['yachttype']}"

    yachtname = f"{row['yachtname']}"
    yachtname = yachtname[0:4] if 'LOON' in yachtname else yachtname

    captain = f"{row['captainname']}".split(" ")
    captain_first_name = captain[0]
    isnan_pre = pd.isnull(df2.iloc[index]['predeliverydateutc'])
    if isnan_pre:
        start = f"{row['startdateutc']}"
    else:
        start = f"{row['predeliverydateutc']}"
    isnan_post = pd.isnull(df2.iloc[index]['postdeliverydateutc'])
    if isnan_post:
        end = f"{row['enddateutc']}"
    else:
        end = f"{row['postdeliverydateutc']}"
    start = start[0:16]
    end = end[0:16]

    for contract_type in df2.iterrows():
        if f"{row['type']}" == 'TRANSPORT':
            doc = docx.Document()
            p = doc.add_paragraph()
            run = p.add_run(f"""Dear {captain_first_name},

I hope you are well.

This email to confirm that we are connected to the AIS of {yachttype} {yachtname} for your transport agreement from {start} to {end}.

Just a reminder that you are registered with FLYN-YACHTING under a PREMIUM contract. This allows you, in addition to the REGULAR service, to benefit from a complete certificate and the possibility of personalized advice during your contract.

Do not hesitate to contact FLYN at any time of the navigation for any advice (in particular the best routes for VAT/70%) or information.
        """)
            run.font.name = 'Arial'
            run.font.size = docx.shared.Pt(12)
            p = doc.add_paragraph()
            run = p.add_run("""VAT""")
            title()
            p = doc.add_paragraph()
            run = p.add_run("""The VAT will be applied only on distances covered within the territorial waters of the country of departure of the trip. Outside the territorial waters, VAT is not applicable.

The applicable VAT rate is that of the country of departure for transport services (10% for navigation in French continental waters, 0% for international navigation). The trip is qualified as "international" if:
        •	(option a) the final port of the considered trip, qualified by a disembarkation and/or embarkation of passengers, is in a country different from the country of departure
        •	(option b) a commercial stopover without disembarkation of passengers, but with purchase of goods and services on board and/or by the passengers, took place in a country different from the country of departure.

Trips qualified as "international" are totally exempt from transport VAT on departure from all countries of the European Union.  There may be reduced transport VAT rates: this is the case for maritime transport in the French territorial waters of Corsica: reduced rate of 2.1% instead of 10%.

                            """)
            main_body_text()
            p = doc.add_paragraph()
            run = p.add_run("""70%""")
            title()
            p = doc.add_paragraph()
            run = p.add_run("""With regard to the 70% high seas navigation criteria, please note that the pre-delivery and post-delivery are each considered as a "fiscal trip" and will count in the final ratio.

Each time you have a stopover involving a change in your passenger list (new embarkation or final disembarkation of at least one passenger) during the cruise, this will generate a new "fiscal trip". It is then important to inform FLYN-YACHTING by sending an email so that we can calculate the final 70% ratio of this cruise. Ordinary passenger movements during transit stops (no change in the passenger list) are not affected by this declaration.

For a fiscal trip to be counted as "high seas", it must include some cruising outside French territorial waters. In order to maximize the high-seas rate of this cruise (high seas trips / total trips), FLYN-YACHTING suggests that you leave French or Monegasque territorial waters at least once per fiscal trip, and spend at least 15 minutes there, when this is acceptable to your client. In most cases, the fuel cost of a return trip outside territorial waters is negligible compared to the fiscal impact of not qualifying for the 70% criteria. The former is at your client’s cost while the latter is at your Owner’s expense.
        """)
            main_body_text()
            p = doc.add_paragraph()
            run = p.add_run("""AIS""")
            title()
            p = doc.add_paragraph()
            run = p.add_run("""In order to provide you these calculations, FLYN-YACHTING will connect to your AIS to register your positions. Sometimes, it is possible to experience a loss of signal from the AIS, which will create a gap in the navigation data of the yacht. For this reason, FLYN-YACHTING recommends you to (a) privilege AIS type A (b) continuously record your positions with your onboard navigation system (from pre to post delivery at 15 minutes of time intervals). This will allow us to recover the missing data in case of signal loss (this service will be subject to an additional charge).
                            """)
            main_body_text()
            p = doc.add_paragraph()
            run = p.add_run("""CONTRACT""")
            title()
            p = doc.add_paragraph()
            run = p.add_run("""In the case of a change to the initial contract (change of start time, end time, port of embarkation, port of disembarkation etc.) please inform us as soon as possible and send us the corresponding addendum.

We remain at your disposal for any further information needed.

Kind regards.

Noemie""")
            main_body_text()
            doc.save("C:\project_data\templates\TRANSPORT_CONTRACT_TEMPLATE_temp.docx")
            a = "C:\project_data\templates\TRANSPORT_CONTRACT_TEMPLATE_temp.docx"
            b = f"C:\project_data\templates\Email_{yachtname}_{id}_TRANSPORT.docx"

        elif f"{row['type']}" == 'CHARTER':
            doc = docx.Document()
            p = doc.add_paragraph()
            run = p.add_run(f"""Dear {captain_first_name},

I hope you are well.

This email to confirm that we are connected to the AIS of {yachttype} {yachtname} for the follow-up of your charter from {start} to {end} to calculate the time spent in EU waters and non-EU waters along your cruise and the 70% high seas navigation count.  Just a reminder that you are registered with FLYN-YACHTING under a PREMIUM contract. This allows you, in addition to the REGULAR service, to benefit from a complete certificate and the possibility of personalized advice during your contract.  Do not hesitate to contact me at any time of the navigation for any advice (in particular the best routes for VAT/70%) or information.
        """)
            run.font.name = 'Arial'
            run.font.size = docx.shared.Pt(12)
            p = doc.add_paragraph()
            run = p.add_run("""VAT""")
            title()
            p = doc.add_paragraph()
            run = p.add_run("""The VAT due by your Client on the Charter contract will only relate to the part of the time spent in EU waters. FLYN-YACHTING will therefore calculate from your AIS positions the time spent in the various territorial and international waters.  FLYN-YACHTING suggests you to cruise as long as possible in non-EU waters in order to benefit from a VAT reduction at the end of the cruise, this of course within the limits of your client's wish to stay or cruise in coastal areas.
                            """)
            main_body_text()
            p = doc.add_paragraph()
            run = p.add_run("""70%""")
            title()
            p = doc.add_paragraph()
            run = p.add_run("""Each time you have a stopover involving a change in your passenger list (new embarkation or final disembarkation of at least one passenger) during the cruise, this will generate a new "fiscal trip". It is then important to inform FLYN-YACHTING by sending us an email so that we can calculate the final 70% ratio of this cruise. Ordinary passenger movements during transit stops (no change in the passenger list) are not affected by this declaration.

For a fiscal trip to be counted as "high seas" from the French point of view, it must include some cruising outside French territorial waters.  As this charter takes place in Italian waters, trips will qualify as "high seas trips" from the French point of view.  With regard to the 70% high seas navigation criteria in Italy, please note that there is only one trip per commercial operation. The disembarkation and embarkation of passengers during the charter are not taken into account.  For this fiscal trip to be counted as "high seas", it must include some cruising outside Italian territorial waters.
        """)
            main_body_text()
            p = doc.add_paragraph()
            run = p.add_run("""AIS""")
            title()
            p = doc.add_paragraph()
            run = p.add_run("""In order to provide you these calculations, FLYN-YACHTING will connect to your AIS to register your positions. Sometimes, it is possible to experience a loss of signal from the AIS, which will create a gap in the navigation data of the yacht. For this reason, FLYN-YACHTING recommends you to (a) privilege AIS type A (b) continuously record your positions with your onboard navigation system (from pre to post delivery at 15 minutes of time intervals).  This will allow us to recover the missing data in case of signal loss.
                            """)
            main_body_text()
            p = doc.add_paragraph()
            run = p.add_run("""CONTRACT""")
            title()
            p = doc.add_paragraph()
            run = p.add_run("""In the case of a change to the initial contract (change of start time, end time, port of embarkation, port of disembarkation etc.) please inform us as soon as possible and send us the corresponding addendum.

We remain at your disposal for any further information needed.

Kind regards.

Noemie""")
            main_body_text()
            doc.save("C:\project_data\templates\CHARTER_CONTRACT_TEMPLATE_temp.docx")
            a = "C:\project_data\templates\CHARTER_CONTRACT_TEMPLATE_temp.docx"
            b = f"C:\project_data\templates\Email_{yachtname}_{id}_CHARTER.docx"

    os.renames(a,b)

outlook = win32.Dispatch('outlook.application')
mail = outlook.CreateItem(0)

path = "C:\project_data\templates\\"
dir_list = os.listdir(path)

for attachment in dir_list:
    path_attach = f'C:\project_data\templates\{attachment}'
    mail.Attachments.Add(path_attach)

mail.Attachments.Add(logo_absolute)
image = mail.Attachments.Add(logo_absolute)

    mail.To = 'operations-team1@example.com'
#   mail.To = 'operations-team1@example.com;operations-team2@example.com;operations-team3@example.com'     #activate when building .exe

mail.Subject = f'{no_pr} PREMIUM CONTRACT(S) STARTING TODAY {today} - automatic email'
mail.HTMLBody += """<div>
                        <h1 style="font-family: 'Calibri'; font-size: 23; font-weight: bold; color: #0A506E;">
                            PREMIUM CONTRACTS STARTING TODAY - please send the email template at least 2 hours before departure.
                        </h1>
                    </div>"""

for contract in range (0,no_pr):
    mail.HTMLBody += f"""<h4 style="font-family: 'Calibri'; font-size: 14; color: #080909;">CUSTOMER: {df2['customerid'][contract]} -
    {df2['producttype'][contract]} | {df2['type'][contract]} |
    {df2['yachttype'][contract]}_{df2['yachtname'][contract]}_{df2['mmsi'][contract]}_{df2['id'][contract]}
    | Pre delivery: {df2['predeliverydateutc'][contract]} | Contract Start: {df2['startdateutc'][contract]} |
    Captain name: {df2['captainname'][contract]}, number: {df2['captaincellular'][contract]}, email: {df2['captainemail'][contract]}</h4>"""
mail.HTMLBody += """<div><i><h9 style="font-family: 'Calibri'; font-size: 13; color: #129090;">*Please modify attached email templates to match desired formatting.</h9></i></div>"""
mail.HTMLBody += """ <div><br><img src="cid:logo-img" width=5%></div>"""
image.PropertyAccessor.SetProperty("http://schemas.microsoft.com/mapi/proptag/0x3712001F", "logo-img")
mail.Send()

for file in dir_list:
    os.remove( f'C:\project_data\templates\{file}')
